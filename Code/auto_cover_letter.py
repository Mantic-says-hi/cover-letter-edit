import os
import sys
import argparse
from docx import Document
from docx2pdf import convert
from docx.shared import Pt


def parse_arguments():
    parser = argparse.ArgumentParser()
    parser.add_argument("position")
    parser.add_argument("company")
    parser.add_argument("applicant")
    parser.add_argument("street")
    parser.add_argument("suburb")
    return parser.parse_args()


def replace_placeholders(document, placeholders):
    for i, paragraph in enumerate(document.paragraphs):
        #Do this to skip last paragraph 
        #Since it deleted my details at the end for some reason without this
        if i == len(document.paragraphs) - 1 and not paragraph.text.strip():
            continue
        for placeholder, value in placeholders.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)
        #Reset font and size back to how I had it since this likes to 
        #Change it to Calibri and size 11
        paragraph.style.font.name = 'Times New Roman'
        paragraph.style.font.size = Pt(12)
            


def generate_cover_letter(template_path, output_path, placeholders):
    try:
        document = Document(template_path)
        replace_placeholders(document, placeholders)
        temp_out = os.path.join(sys.path[0], "tempOut.docx")
        document.save(temp_out)
        convert(temp_out, output_path)
        os.remove(temp_out)
    except Exception as e:
        print(f"An error occurred: {e}")


if __name__ == "__main__":
    args = parse_arguments()
    template_path = os.environ.get('COVER_LETTER_TEMPLATE_PATH')
    output_filename = f"{args.applicant} Cover Letter for {args.company}.pdf"
    output_path = os.path.join(os.path.abspath(os.getcwd()), output_filename)
    
    placeholders = {
        "[POSITION]": args.position.upper(),
        "[JOB TITLE]": args.position,
        "[COMPANY NAME]": args.company,
        "[APPLICANT]": args.applicant,
        "[STREET]": args.street,
        "[SUBURB]": args.suburb,
    }
    generate_cover_letter(template_path, output_path, placeholders)